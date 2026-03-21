"""Universal document loader — extracts text from PDF, Word, Excel, and text files.

Chunks documents into manageable pieces for ChromaDB embedding.
Supports any document type: bank SOPs, research papers, company guidelines, etc.
"""

import hashlib
import io
import uuid
from pathlib import Path
from typing import BinaryIO

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from openpyxl import load_workbook


# ---------------------------------------------------------------------------
# Text extraction per format
# ---------------------------------------------------------------------------

def extract_pdf(source: str | Path | BinaryIO) -> str:
    """Extract all text from a PDF file."""
    if isinstance(source, (str, Path)):
        doc = fitz.open(str(source))
    else:
        data = source.read()
        doc = fitz.open(stream=data, filetype="pdf")

    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n\n".join(pages)


def extract_docx(source: str | Path | BinaryIO) -> str:
    """Extract all text from a Word .docx file."""
    if isinstance(source, (str, Path)):
        doc = DocxDocument(str(source))
    else:
        doc = DocxDocument(source)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def extract_excel(source: str | Path | BinaryIO) -> str:
    """Extract text from all sheets of an Excel file."""
    if isinstance(source, (str, Path)):
        wb = load_workbook(str(source), read_only=True, data_only=True)
    else:
        wb = load_workbook(source, read_only=True, data_only=True)

    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"--- Sheet: {sheet_name} ---")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line = " | ".join(cells).strip()
            if line and line != "| " * len(cells):
                parts.append(line)
    wb.close()
    return "\n".join(parts)


def extract_text(source: str | Path | BinaryIO) -> str:
    """Extract text from a plain text / markdown file."""
    if isinstance(source, (str, Path)):
        return Path(source).read_text(encoding="utf-8")
    else:
        raw = source.read()
        if isinstance(raw, bytes):
            return raw.decode("utf-8", errors="replace")
        return raw


# ---------------------------------------------------------------------------
# Format detection
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".txt", ".md", ".csv", ".json"}

_EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".xlsx": extract_excel,
    ".xls": extract_excel,
    ".txt": extract_text,
    ".md": extract_text,
    ".csv": extract_text,
    ".json": extract_text,
}


def detect_and_extract(filename: str, source: str | Path | BinaryIO) -> str:
    """Detect format from filename and extract text."""
    ext = Path(filename).suffix.lower()
    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        raise ValueError(f"Unsupported file format: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}")
    return extractor(source)


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
    """Split text into overlapping chunks for embedding.

    Args:
        text: Full document text.
        chunk_size: Target characters per chunk.
        overlap: Characters of overlap between adjacent chunks.

    Returns:
        List of text chunks.
    """
    if not text.strip():
        return []

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]

        # Try to break at a paragraph or sentence boundary
        if end < len(text):
            # Look for last paragraph break
            last_para = chunk.rfind("\n\n")
            if last_para > chunk_size // 2:
                end = start + last_para + 2
                chunk = text[start:end]
            else:
                # Look for last sentence break
                last_period = chunk.rfind(". ")
                if last_period > chunk_size // 2:
                    end = start + last_period + 2
                    chunk = text[start:end]

        chunk = chunk.strip()
        if chunk:
            chunks.append(chunk)

        start = end - overlap
        if start <= 0 and end >= len(text):
            break

    return chunks


def file_content_hash(content: bytes) -> str:
    """Generate a SHA-256 hash of file content for deduplication."""
    return hashlib.sha256(content).hexdigest()[:16]


def generate_chunk_id(filename: str, chunk_index: int, content_hash: str) -> str:
    """Generate a deterministic chunk ID."""
    return f"{Path(filename).stem}_{content_hash}_{chunk_index:04d}"
